from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
OUTPUT_DIR = ROOT / "docs" / "cv"
OUTPUT_PATH = OUTPUT_DIR / "VU_THE_LUC_CV_BTC.docx"
LOGO_PATH = ROOT / "docs" / "demo" / "brand" / "png" / "logo" / "nera-symbol-light.png"

# compact_reference_guide with a named one-page CV override.
FONT = "Arial"
FOREST = "14372F"
DEEP_FOREST = "09241E"
SAGE = "DDE9E1"
CANVAS = "F7F4ED"
CORAL = "E26D4F"
INK = "10261F"
MUTED = "60716A"
WHITE = "FFFFFF"
LIGHT_BORDER = "DCE3DE"

PAGE_WIDTH_DXA = 12240
CONTENT_WIDTH_DXA = 10224
HEADER_LEFT_DXA = 8800
HEADER_RIGHT_DXA = CONTENT_WIDTH_DXA - HEADER_LEFT_DXA
METRIC_WIDTHS = [3408, 3408, 3408]


def rgb(hex_color: str) -> RGBColor:
    return RGBColor.from_string(hex_color)


def set_run_font(run, *, size=None, color=INK, bold=None, italic=None, name=FONT):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = rgb(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_paragraph_spacing(paragraph, *, before=0, after=0, line=1.0, keep_next=False):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line
    fmt.keep_with_next = keep_next


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin_name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin_name}"))
        if node is None:
            node = OxmlElement(f"w:{margin_name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color=LIGHT_BORDER, size=4, val="single"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), val)
        if val != "nil":
            tag.set(qn("w:sz"), str(size))
            tag.set(qn("w:color"), color)


def set_table_geometry(table, widths, indent=0):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr

    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")

    layout = tbl_pr.first_child_found_in("w:tblLayout")
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            width = widths[min(idx, len(widths) - 1)]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")


def add_hyperlink(paragraph, text: str, url: str, *, color=FOREST, bold=False, size=9.2):
    part = paragraph.part
    rel_id = part.relate_to(url, RT.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    r_fonts = OxmlElement("w:rFonts")
    r_fonts.set(qn("w:ascii"), FONT)
    r_fonts.set(qn("w:hAnsi"), FONT)
    color_node = OxmlElement("w:color")
    color_node.set(qn("w:val"), color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "none")
    size_node = OxmlElement("w:sz")
    size_node.set(qn("w:val"), str(int(size * 2)))
    r_pr.extend([r_fonts, color_node, underline, size_node])
    if bold:
        r_pr.append(OxmlElement("w:b"))
    run.append(r_pr)
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_bullet_numbering(document) -> int:
    numbering = document.part.numbering_part.element
    abstract_ids = [int(node.get(qn("w:abstractNumId"))) for node in numbering.findall(qn("w:abstractNum"))]
    num_ids = [int(node.get(qn("w:numId"))) for node in numbering.findall(qn("w:num"))]
    abstract_id = max(abstract_ids, default=0) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "bullet")
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "•")
    lvl_jc = OxmlElement("w:lvlJc")
    lvl_jc.set(qn("w:val"), "left")
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "540")
    ind.set(qn("w:hanging"), "270")
    p_pr.extend([tabs, ind])
    r_pr = OxmlElement("w:rPr")
    color_node = OxmlElement("w:color")
    color_node.set(qn("w:val"), CORAL)
    r_pr.append(color_node)
    lvl.extend([start, num_fmt, lvl_text, lvl_jc, p_pr, r_pr])
    abstract.append(lvl)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def apply_bullet(paragraph, num_id: int):
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.append(num_pr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id_node = OxmlElement("w:numId")
    num_id_node.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, num_id_node])


def add_section_heading(document, text: str):
    paragraph = document.add_paragraph(style="Heading 1")
    set_paragraph_spacing(paragraph, before=7, after=3, line=1.0, keep_next=True)
    run = paragraph.add_run(text.upper())
    set_run_font(run, size=10.5, color=CORAL, bold=True)
    run.font.letter_spacing = Pt(1.2) if hasattr(run.font, "letter_spacing") else None
    return paragraph


def add_labeled_line(document, label: str, value: str):
    p = document.add_paragraph()
    set_paragraph_spacing(p, after=2.2, line=1.05)
    label_run = p.add_run(label + ": ")
    set_run_font(label_run, size=9.2, color=FOREST, bold=True)
    value_run = p.add_run(value)
    set_run_font(value_run, size=9.2, color=INK)
    return p


def build_document():
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    document = Document()
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.58)
    section.bottom_margin = Inches(0.52)
    section.left_margin = Inches(0.70)
    section.right_margin = Inches(0.70)
    section.header_distance = Inches(0.30)
    section.footer_distance = Inches(0.30)

    document.core_properties.title = "CV Vũ Thế Lực - BTC"
    document.core_properties.subject = "AI Product, Business Analyst, AI Quality & Delivery"
    document.core_properties.author = "Vũ Thế Lực"
    document.core_properties.keywords = "AI Product, Business Analyst, Nera, P-046, AI20K"

    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    normal.font.size = Pt(9.4)
    normal.font.color.rgb = rgb(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(3)
    normal.paragraph_format.line_spacing = 1.05

    for style_name, size, color in (("Heading 1", 10.5, CORAL), ("Heading 2", 11.5, FOREST), ("Heading 3", 10.0, FOREST)):
        style = styles[style_name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:ascii"), FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = rgb(color)

    bullet_num_id = add_bullet_numbering(document)

    header = document.add_table(rows=1, cols=2)
    set_table_geometry(header, [HEADER_LEFT_DXA, HEADER_RIGHT_DXA], indent=0)
    set_table_borders(header, val="nil")
    left, right = header.rows[0].cells
    set_cell_margins(left, top=0, start=0, bottom=0, end=120)
    set_cell_margins(right, top=0, start=0, bottom=0, end=0)
    left.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    right.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    p = left.paragraphs[0]
    set_paragraph_spacing(p, after=1, line=1.0)
    name_run = p.add_run("VŨ THẾ LỰC")
    set_run_font(name_run, size=27, color=DEEP_FOREST, bold=True)

    p = left.add_paragraph()
    set_paragraph_spacing(p, after=5, line=1.0)
    role_run = p.add_run("AI PRODUCT  •  BUSINESS ANALYST  •  AI QUALITY & DELIVERY")
    set_run_font(role_run, size=10.0, color=CORAL, bold=True)

    p = left.add_paragraph()
    set_paragraph_spacing(p, after=0, line=1.0)
    set_run_font(p.add_run("2A202602008  •  "), size=9.0, color=MUTED)
    add_hyperlink(p, "lucvu0728@gmail.com", "mailto:lucvu0728@gmail.com", color=FOREST, size=9.0)
    set_run_font(p.add_run("  •  "), size=9.0, color=MUTED)
    add_hyperlink(p, "github.com/Lucvuu", "https://github.com/Lucvuu", color=FOREST, size=9.0)

    logo_p = right.paragraphs[0]
    logo_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_paragraph_spacing(logo_p, after=0)
    if LOGO_PATH.exists():
        logo_run = logo_p.add_run()
        inline = logo_run.add_picture(str(LOGO_PATH), width=Inches(0.75))
        doc_pr = inline._inline.docPr
        doc_pr.set("descr", "Biểu tượng Nera - dự án P-046")

    summary = document.add_paragraph()
    set_paragraph_spacing(summary, before=5, after=5, line=1.10)
    summary.paragraph_format.left_indent = Inches(0.02)
    summary.paragraph_format.right_indent = Inches(0.02)
    s1 = summary.add_run("AI Product/BA của đội P-046, kết nối phản hồi mentor với định hướng sản phẩm, acceptance criteria, kiểm thử tích hợp và bằng chứng báo cáo. ")
    set_run_font(s1, size=9.7, color=INK)
    s2 = summary.add_run("Làm việc hiệu quả ở giao điểm sản phẩm - kỹ thuật, ưu tiên AI có grounding, khả năng giải thích và con người xác nhận hành động quan trọng.")
    set_run_font(s2, size=9.7, color=INK)

    metrics = document.add_table(rows=1, cols=3)
    set_table_geometry(metrics, METRIC_WIDTHS, indent=0)
    set_table_borders(metrics, val="nil")
    metric_data = [
        ("MVP LIVE", "nerahome.space"),
        ("QUALITY GATE", "53 tests pass"),
        ("CONTRIBUTION", "39 commit không tính merge"),
    ]
    for idx, (label, value) in enumerate(metric_data):
        cell = metrics.rows[0].cells[idx]
        set_cell_shading(cell, FOREST if idx != 1 else DEEP_FOREST)
        set_cell_margins(cell, top=105, start=150, bottom=105, end=150)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph_spacing(p, after=2, line=1.0)
        set_run_font(p.add_run(label), size=7.9, color=SAGE, bold=True)
        p = cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph_spacing(p, after=0, line=1.0)
        set_run_font(p.add_run(value), size=10.2, color=WHITE, bold=True)

    add_section_heading(document, "Kinh nghiệm dự án tiêu biểu")
    p = document.add_paragraph(style="Heading 2")
    set_paragraph_spacing(p, after=1, line=1.0, keep_next=True)
    set_run_font(p.add_run("NERA - AI HOME COMPANION"), size=12.2, color=FOREST, bold=True)
    set_run_font(p.add_run("   |   046LTD / P-046   |   08/2026"), size=8.8, color=MUTED, bold=True)

    p = document.add_paragraph()
    set_paragraph_spacing(p, after=2.5, line=1.0, keep_next=True)
    set_run_font(p.add_run("Vai trò: "), size=9.2, color=CORAL, bold=True)
    set_run_font(p.add_run("AI Product, Business Analyst, Integration QA & Reporting"), size=9.2, color=INK, bold=True)

    project_bullets = [
        "Chuyển phản hồi mentor thành định hướng V2 và hệ thống tài liệu điều hành: Product Brief, MVP Scope, Product Outcomes, User Journey, decision/risk/action logs và acceptance criteria.",
        "Review giao điểm đăng nhập, session ownership, memory, tìm kiếm và booking; harden quyền sở hữu phiên, property matching và bổ sung regression coverage cho các ca chọn nhầm bất động sản.",
        "Chuẩn bị bằng chứng Gate 2 và Demo Day Phase 1: kiến trúc, eval evidence, báo cáo, demo script, submission pack, slide, nguồn NotebookLM và bộ nhận diện Nera; kiểm soát claim theo trạng thái thật.",
        "Đóng góp kỹ thuật qua AI Elements spike, test batch, trace/observability, bảo vệ AI log key, sửa lỗi agent/frontend và loại bỏ 1.814 dòng dead code đã xác nhận.",
    ]
    for text in project_bullets:
        p = document.add_paragraph()
        apply_bullet(p, bullet_num_id)
        set_paragraph_spacing(p, after=2.1, line=1.05)
        set_run_font(p.add_run(text), size=9.15, color=INK)

    links = document.add_paragraph()
    set_paragraph_spacing(links, before=1, after=0, line=1.0)
    set_run_font(links.add_run("Sản phẩm: "), size=8.8, color=MUTED, bold=True)
    add_hyperlink(links, "www.nerahome.space", "https://www.nerahome.space/", color=FOREST, bold=True, size=8.8)
    set_run_font(links.add_run("   •   Repository: "), size=8.8, color=MUTED, bold=True)
    add_hyperlink(links, "AI20K-Build-Phase-Cohort-3/P-046", "https://github.com/AI20K-Build-Phase-Cohort-3/P-046", color=FOREST, bold=True, size=8.8)

    add_section_heading(document, "Năng lực cốt lõi")
    add_labeled_line(document, "Product & BA", "Product discovery, phân tích yêu cầu, product brief/PRD, user journey, acceptance criteria, scope và decision/risk management")
    add_labeled_line(document, "AI Quality", "Agent testing, eval evidence, grounding, human-in-the-loop, trace/observability, regression test và UAT")
    add_labeled_line(document, "Technical collaboration", "LangGraph, FastAPI, PostgreSQL, Redis, Next.js, Git/GitHub, Figma, Vercel và Render ở mức làm việc")

    add_section_heading(document, "Chương trình & hoạt động")
    p = document.add_paragraph()
    set_paragraph_spacing(p, after=1.5, line=1.05)
    set_run_font(p.add_run("AI20K BUILD PHASE COHORT 3"), size=9.5, color=FOREST, bold=True)
    set_run_font(p.add_run("   |   Team 046LTD   |   2026"), size=8.8, color=MUTED, bold=True)
    p = document.add_paragraph()
    set_paragraph_spacing(p, after=0, line=1.05)
    set_run_font(p.add_run("Thực hành dự án theo nhóm qua lab/workshop; phụ trách product alignment, BA, kiểm soát chất lượng và báo cáo mentor/Demo Day."), size=9.15, color=INK)

    footer = section.footer
    footer_p = footer.paragraphs[0]
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(footer_p, after=0, line=1.0)
    set_run_font(footer_p.add_run("CV nộp Ban Tổ Chức  •  P-046 / Nera  •  Cập nhật 25/08/2026"), size=7.5, color=MUTED)

    document.save(OUTPUT_PATH)
    return OUTPUT_PATH


if __name__ == "__main__":
    path = build_document()
    print(path)
