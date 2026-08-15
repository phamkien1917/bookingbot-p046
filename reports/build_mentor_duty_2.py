from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path(__file__).with_name("Bao_cao_tien_do_Mentor_Duty_2_XHome_VisitOps.docx")

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
NAVY = "17365D"
MUTED = "667085"
LIGHT_GRAY = "F2F4F7"
PALE_BLUE = "E8EEF5"
PALE_GREEN = "EAF6EC"
PALE_GOLD = "FFF4D6"
PALE_RED = "FDECEC"
WHITE = "FFFFFF"
BLACK = "000000"
BORDER = "C9D2DC"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=120, bottom=90, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color=BORDER, size=6):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = qn(f"w:{edge}")
        element = borders.find(tag)
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), str(size))
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_table_geometry(table, widths_dxa, indent_dxa=120):
    total = sum(widths_dxa)
    table.autofit = False
    tbl_pr = table._tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths_dxa[idx]))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Inches(widths_dxa[idx] / 1440)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def prevent_row_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    cant_split.set(qn("w:val"), "true")
    tr_pr.append(cant_split)


def set_repeat_table_header(row):
    repeat_table_header(row)


def set_run(run, size=11, bold=False, color=BLACK, italic=False, font="Calibri"):
    run.font.name = font
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), font)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), font)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)
    return run


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Trang ")
    set_run(run, size=9, color=MUTED)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])


def add_rule(paragraph, color=BLUE, size=12):
    p_pr = paragraph._p.get_or_add_pPr()
    borders = p_pr.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        p_pr.append(borders)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), "6")
    bottom.set(qn("w:color"), color)
    borders.append(bottom)


def apply_num_id(paragraph, num_id, level=0):
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), str(level))
    num_id_node = OxmlElement("w:numId")
    num_id_node.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, num_id_node])
    p_pr.append(num_pr)


def add_bullet(doc, text, level=0, color=BLACK):
    if not hasattr(doc, "_xhome_bullet_num_id"):
        doc._xhome_bullet_num_id = create_bullet_numbering_id(doc)
    p = doc.add_paragraph()
    apply_num_id(p, doc._xhome_bullet_num_id, level=0)
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.line_spacing = 1.167
    p.paragraph_format.keep_together = True
    set_run(p.add_run(text), size=11, color=color)
    return p


def create_numbering_id(doc):
    numbering = doc.part.numbering_part.element
    abstract_ids = [int(n.get(qn("w:abstractNumId"))) for n in numbering.xpath("./w:abstractNum")]
    num_ids = [int(n.get(qn("w:numId"))) for n in numbering.xpath("./w:num")]
    abstract_id = max(abstract_ids, default=0) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "decimal")
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "%1.")
    suff = OxmlElement("w:suff")
    suff.set(qn("w:val"), "tab")
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "720")
    tabs.append(tab)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "720")
    ind.set(qn("w:hanging"), "360")
    p_pr.extend([tabs, ind])
    lvl.extend([start, num_fmt, lvl_text, suff, p_pr])
    abstract.append(lvl)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    override = OxmlElement("w:lvlOverride")
    override.set(qn("w:ilvl"), "0")
    start_override = OxmlElement("w:startOverride")
    start_override.set(qn("w:val"), "1")
    override.append(start_override)
    num.append(override)
    numbering.append(num)
    return num_id


def create_bullet_numbering_id(doc):
    numbering = doc.part.numbering_part.element
    abstract_ids = [int(n.get(qn("w:abstractNumId"))) for n in numbering.xpath("./w:abstractNum")]
    num_ids = [int(n.get(qn("w:numId"))) for n in numbering.xpath("./w:num")]
    abstract_id = max(abstract_ids, default=0) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "bullet")
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "•")
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "720")
    tabs.append(tab)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "720")
    ind.set(qn("w:hanging"), "360")
    p_pr.extend([tabs, ind])
    r_pr = OxmlElement("w:rPr")
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), "Calibri")
    fonts.set(qn("w:hAnsi"), "Calibri")
    r_pr.append(fonts)
    lvl.extend([start, num_fmt, lvl_text, p_pr, r_pr])
    abstract.append(lvl)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def add_number(doc, text, num_id):
    p = doc.add_paragraph()
    apply_num_id(p, num_id)
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.line_spacing = 1.167
    p.paragraph_format.keep_together = True
    set_run(p.add_run(text), size=11)
    return p


def add_body(doc, text, bold_prefix=None, color=BLACK):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.10
    if bold_prefix and text.startswith(bold_prefix):
        set_run(p.add_run(bold_prefix), bold=True, color=color)
        set_run(p.add_run(text[len(bold_prefix):]), color=color)
    else:
        set_run(p.add_run(text), color=color)
    return p


def add_callout(doc, label, text, fill=PALE_BLUE, label_color=DARK_BLUE):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    set_table_borders(table, color=fill, size=0)
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.10
    set_run(p.add_run(f"{label}: "), bold=True, color=label_color)
    set_run(p.add_run(text), color=BLACK)
    after = doc.add_paragraph()
    after.paragraph_format.space_after = Pt(2)
    return table


def add_status_table(doc, headers, rows, widths_dxa, status_col=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_geometry(table, widths_dxa)
    set_table_borders(table)
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    prevent_row_split(hdr)
    for idx, header in enumerate(headers):
        set_cell_shading(hdr.cells[idx], NAVY)
        p = hdr.cells[idx].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        set_run(p.add_run(header), size=9.5, bold=True, color=WHITE)

    for row_data in rows:
        row = table.add_row()
        prevent_row_split(row)
        for idx, value in enumerate(row_data):
            cell = row.cells[idx]
            if status_col is not None and idx == status_col:
                status = str(value)
                if status.startswith("Hoàn thành"):
                    set_cell_shading(cell, PALE_GREEN)
                elif status.startswith("Đang"):
                    set_cell_shading(cell, PALE_GOLD)
                elif status.startswith("Chưa"):
                    set_cell_shading(cell, PALE_RED)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.05
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if idx in (0, status_col) else WD_ALIGN_PARAGRAPH.LEFT
            set_run(p.add_run(str(value)), size=9.2, color=BLACK)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(2)
    return table


def keep_with_next(paragraph):
    paragraph.paragraph_format.keep_with_next = True


doc = Document()
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1)
section.right_margin = Inches(1)
section.header_distance = Inches(0.492)
section.footer_distance = Inches(0.492)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Calibri"
normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
normal.font.size = Pt(11)
normal.paragraph_format.space_before = Pt(0)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.10

for name, size, color, before, after in (
    ("Heading 1", 16, BLUE, 16, 8),
    ("Heading 2", 13, BLUE, 12, 6),
    ("Heading 3", 12, DARK_BLUE, 8, 4),
):
    style = styles[name]
    style.font.name = "Calibri"
    style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    style.font.size = Pt(size)
    style.font.bold = True
    style.font.color.rgb = RGBColor.from_string(color)
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)
    style.paragraph_format.keep_with_next = True

for list_name in ("List Bullet", "List Bullet 2", "List Number"):
    st = styles[list_name]
    st.font.name = "Calibri"
    st._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    st._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    st.font.size = Pt(11)
    st.paragraph_format.left_indent = Inches(0.5 if list_name != "List Bullet 2" else 0.75)
    st.paragraph_format.first_line_indent = Inches(-0.25)
    st.paragraph_format.space_after = Pt(8)
    st.paragraph_format.line_spacing = 1.167

header = section.header
hp = header.paragraphs[0]
hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
set_run(hp.add_run("AI20K | P-046 | XHome VisitOps"), size=9, bold=True, color=MUTED)

footer = section.footer
fp = footer.paragraphs[0]
add_page_number(fp)

# First-page memo masthead.
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(2)
set_run(p.add_run("BÁO CÁO TIẾN ĐỘ"), size=24, bold=True, color=NAVY)

p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(14)
set_run(p.add_run("Mentor Duty lần 2 - MVP hệ thống đặt lịch xem bất động sản O2O"), size=14, bold=True, color=DARK_BLUE)

metadata = [
    ("Dự án", "XHome VisitOps - BĐS Kinh doanh O2O"),
    ("Nhóm", "P-046"),
    ("Mentor", "Đặng Hải Lộc"),
    ("Ngày báo cáo", "01/08/2026"),
    ("Giai đoạn", "Chốt phạm vi MVP và hoàn thiện nền tảng dữ liệu"),
    ("Trạng thái", "Database foundation hoàn thành; application layer đang khởi động"),
]
for label, value in metadata:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.0
    set_run(p.add_run(f"{label}: "), bold=True, color=BLACK)
    set_run(p.add_run(value), color=BLACK)

rule = doc.add_paragraph()
rule.paragraph_format.space_after = Pt(10)
add_rule(rule)

add_callout(
    doc,
    "Thông điệp chính",
    "Nhóm đã chốt được bài toán, luồng nghiệp vụ và schema PostgreSQL 24 bảng chạy trên Neon. Mục tiêu tiếp theo là thu hẹp MVP thành một lát cắt hoàn chỉnh: khách yêu cầu xem nhà -> agent đề xuất giờ -> sale duyệt -> hệ thống tạo lịch và soft-hold an toàn.",
    fill=PALE_BLUE,
)

h = doc.add_heading("1. Tổng quan bài toán", level=1)
keep_with_next(h)
add_body(
    doc,
    "Hiện nay việc hẹn khách xem nhà mẫu hoặc căn thực tế phải phối hợp giữa khách, sale, xe đưa đón, trạng thái căn và phòng chờ. Điều phối thủ công qua chat dễ gây trùng lịch, phản hồi chậm và bỏ lỡ khách tiềm năng."
)
add_body(
    doc,
    "Giải pháp đề xuất là AI Agent hiểu yêu cầu bằng ngôn ngữ tự nhiên, kiểm tra dữ liệu lịch và căn, đề xuất khung giờ, chuyển sale xác nhận (HITL), sau đó tạo lịch và giữ căn tạm thời. Các thao tác quan trọng phải có transaction và ràng buộc chống double-booking."
)

h = doc.add_heading("2. MVP là gì và MVP của nhóm gồm những gì?", level=1)
keep_with_next(h)
add_body(
    doc,
    "MVP (Minimum Viable Product) là phiên bản nhỏ nhất nhưng chạy xuyên suốt được một nhu cầu thực tế và đủ để kiểm chứng giả thuyết sản phẩm. MVP không phải là bản làm sơ sài; nhóm cần ưu tiên một luồng chính hoạt động ổn định trước khi mở rộng."
)

h = doc.add_heading("Luồng demo MVP đề xuất", level=2)
keep_with_next(h)
for item in (
    "Khách đăng nhập và mô tả nhu cầu xem bất động sản bằng tiếng Việt tự nhiên.",
    "Agent hỏi thêm thông tin còn thiếu: căn/dự án, ngày mong muốn, khoảng giờ, điểm đón và số người.",
    "Backend kiểm tra trạng thái căn, lịch bận của sale và khung giờ có thể phục vụ.",
    "Agent trả về 2-3 phương án; khách chọn một phương án.",
    "Sale nhận yêu cầu và bấm duyệt hoặc từ chối (HITL).",
    "Nếu được duyệt, hệ thống tạo appointment, soft-hold căn và thông báo xác nhận trong cùng một transaction.",
    "Khách hoặc sale có thể xem trạng thái, dời hoặc hủy lịch theo quy tắc hệ thống.",
):
    add_bullet(doc, item)

h = doc.add_heading("Phạm vi MVP và phần để sau", level=2)
keep_with_next(h)
add_status_table(
    doc,
    ["Nhóm chức năng", "Trong MVP", "Sau MVP"],
    [
        ("Tài khoản và vai trò", "Khách, sale/điều phối", "SSO, phân quyền chi tiết"),
        ("Hội thoại AI", "Thu thập yêu cầu và gọi tool", "Memory dài hạn, cá nhân hóa sâu"),
        ("Đặt lịch", "Đề xuất giờ, HITL, xác nhận", "Tự tối ưu và dời lịch hàng loạt"),
        ("Tồn kho căn", "Kiểm tra trạng thái, soft-hold", "Đồng bộ nhiều hệ thống inventory"),
        ("Điều phối", "Kiểm tra sale/căn/xe cơ bản", "Tối ưu lộ trình nhiều khách"),
        ("Bản đồ", "Lưu địa chỉ và tọa độ", "Tiện ích xung quanh, routing thực tế"),
        ("Đánh giá", "Log sự kiện cơ bản", "Dashboard success rate/no-show"),
    ],
    [2500, 3300, 3560],
)

h = doc.add_heading("3. Tiến độ thực tế đến Mentor Duty lần 2", level=1)
keep_with_next(h)
add_callout(
    doc,
    "Cách đọc trạng thái",
    "'Hoàn thành' chỉ dùng khi đã có artefact kiểm chứng được trong repository hoặc trên Neon. Code template không được tính là tính năng nghiệp vụ hoàn chỉnh.",
    fill=LIGHT_GRAY,
    label_color=MUTED,
)

progress_rows = [
    ("1", "Phân tích bài toán và ràng buộc", "Hoàn thành", "Luồng booking, HITL, soft-hold, chống trùng lịch đã được mô hình hóa."),
    ("2", "Thiết kế PostgreSQL", "Hoàn thành", "24 bảng: 17 bảng lõi MVP + 7 bảng mở rộng; có PK/FK, index, trigger và exclusion constraint."),
    ("3", "Database cloud", "Hoàn thành", "Schema đã chạy trên Neon; đã quan sát đủ bảng và quan hệ bằng DBeaver ERD."),
    ("4", "Dữ liệu mẫu và kiểm thử DB", "Hoàn thành", "Có 002_seed.sql và 003_smoke_test.sql; cần chạy lại trước mỗi buổi demo chính thức."),
    ("5", "FastAPI backend", "Đang thực hiện", "Có skeleton /health, /api/v1/chat và /status; chưa có API nghiệp vụ booking/inventory/HITL."),
    ("6", "LangGraph agent", "Đang thực hiện", "Có graph mẫu analyze -> respond; chưa có tool kiểm tra lịch/căn hoặc luồng phê duyệt."),
    ("7", "Frontend Next.js", "Chưa bắt đầu", "Chưa thấy mã frontend trong repository hiện tại."),
    ("8", "Authentication", "Chưa bắt đầu", "Chưa có JWT/session và màn hình đăng nhập theo vai trò."),
    ("9", "Calendar/notification", "Chưa bắt đầu", "Chưa tích hợp Google/Outlook Calendar và email/SMS."),
    ("10", "Deploy end-to-end", "Chưa bắt đầu", "Database đã ở cloud nhưng frontend/backend chưa có live URL."),
]
add_status_table(
    doc,
    ["#", "Hạng mục", "Trạng thái", "Bằng chứng / ghi chú"],
    progress_rows,
    [520, 2300, 1550, 4990],
    status_col=2,
)

h = doc.add_heading("Artefact hiện có", level=2)
keep_with_next(h)
for item in (
    "database/001_schema.sql - cấu trúc database và ràng buộc nghiệp vụ.",
    "database/002_seed.sql - dữ liệu mẫu phục vụ demo.",
    "database/003_smoke_test.sql - kiểm tra nhanh các luồng/ràng buộc quan trọng.",
    "docs/database-design.md - mô tả ERD, booking state, HITL, soft-hold và eval.",
    "src/main.py, src/api/routes.py - FastAPI skeleton.",
    "src/agents/graph.py - LangGraph skeleton; chưa phải agent nghiệp vụ hoàn chỉnh.",
):
    add_bullet(doc, item)

h = doc.add_heading("4. Kiến trúc MVP mục tiêu", level=1)
keep_with_next(h)
architecture_rows = [
    ("1. Next.js", "Giao diện khách và sale", "Đăng nhập, chat, chọn slot, duyệt lịch, xem trạng thái"),
    ("2. FastAPI", "API và nghiệp vụ", "Auth, validation, transaction, error handling, WebSocket khi cần"),
    ("3. LangGraph", "Điều phối hội thoại/tool", "Collect requirements -> check availability -> propose -> request approval"),
    ("4. Tool layer", "Tác vụ có kiểm soát", "Inventory, sale calendar, appointment, soft-hold, notification"),
    ("5. Neon PostgreSQL", "Nguồn dữ liệu chuẩn", "Users, properties, requests, approvals, appointments, holds, audit"),
]
add_status_table(
    doc,
    ["Lớp", "Vai trò", "Đầu ra chính"],
    architecture_rows,
    [1800, 2500, 5060],
)
add_callout(
    doc,
    "Nguyên tắc an toàn",
    "LLM không được tự chốt lịch. Agent chỉ đề xuất và gọi tool; sale phải phê duyệt. Transaction database là lớp bảo vệ cuối để chống hai yêu cầu đồng thời giữ cùng sale/căn/xe.",
    fill=PALE_GOLD,
    label_color="7A5A00",
)

h = doc.add_heading("5. Tiêu chí nghiệm thu MVP", level=1)
keep_with_next(h)
criteria = [
    ("AC-01", "Đăng nhập", "Khách và sale truy cập đúng màn hình theo vai trò."),
    ("AC-02", "Hiểu yêu cầu", "Agent trích xuất được căn/dự án, ngày, khoảng giờ và điểm đón từ hội thoại."),
    ("AC-03", "Đề xuất khả dụng", "Chỉ hiển thị slot không xung đột với căn và sale."),
    ("AC-04", "HITL", "Không thể tạo appointment CONFIRMED nếu chưa có approval APPROVED hợp lệ."),
    ("AC-05", "Chống double-booking", "Hai request chồng thời gian cho cùng sale/căn/xe không thể cùng thành công."),
    ("AC-06", "Soft-hold", "Hold có thời điểm hết hạn, giới hạn gia hạn và tự chuyển trạng thái khi quá hạn."),
    ("AC-07", "Xử lý lỗi", "API calendar/inventory lỗi thì không chốt mù; trả trạng thái chờ và cảnh báo điều phối."),
    ("AC-08", "Demo end-to-end", "Một kịch bản từ chat đến appointment + hold chạy được trên live web."),
]
add_status_table(doc, ["ID", "Tiêu chí", "Điều kiện đạt"], criteria, [900, 2400, 6060])

h = doc.add_heading("6. Kế hoạch sprint tiếp theo", level=1)
keep_with_next(h)
add_body(
    doc,
    "Mục tiêu sprint: hoàn thành một vertical slice có thể demo, ưu tiên luồng happy path trước, sau đó mới bổ sung dời/hủy và các tích hợp nâng cao."
)

plan_rows = [
    ("P0", "Backend/Data", "Kết nối FastAPI với Neon; tạo repository/service cho properties, availability, approvals và appointments.", "API nghiệp vụ + transaction booking", "Ngày 1-2"),
    ("P0", "Agent", "Thay graph mẫu bằng state và node nghiệp vụ; tạo tools check_property, check_slots, request_approval.", "Agent gọi tool có kiểm soát", "Ngày 2-3"),
    ("P0", "Frontend", "Tạo Next.js tối thiểu: đăng nhập giả lập theo role, chat khách, màn sale duyệt.", "UI demo 2 vai trò", "Ngày 2-4"),
    ("P0", "Integration", "Nối khách chọn slot -> sale duyệt -> appointment + hold -> xác nhận.", "Happy path end-to-end", "Ngày 4-5"),
    ("P1", "Quality", "Test xung đột lịch, approval hết hạn, hold hết hạn, API tool timeout.", "Test evidence", "Ngày 5-6"),
    ("P1", "Deploy", "Deploy backend/frontend; cấu hình DATABASE_URL và secrets ngoài source code.", "Live URL", "Ngày 6"),
    ("P2", "Polish", "Cải thiện UI, loading/error states, thêm hủy/dời lịch nếu còn thời gian.", "Demo ổn định", "Ngày 7"),
]
add_status_table(
    doc,
    ["Ưu tiên", "Stream", "Công việc", "Đầu ra", "Mốc"],
    plan_rows,
    [900, 1350, 3620, 2260, 1230],
)

h = doc.add_heading("Phân công cần chốt ngay sau buổi họp", level=2)
keep_with_next(h)
for item in (
    "Database/Data owner: Phạm Kiên - quản lý schema, Neon, migration và transaction booking.",
    "Backend owner: chốt một người chịu trách nhiệm API/service và kết nối Neon.",
    "Agent owner: chốt một người chịu trách nhiệm LangGraph state, nodes và tools.",
    "Frontend owner: chốt một người chịu trách nhiệm Next.js cho khách và sale.",
    "Integration/QA owner: chốt một người giữ kịch bản demo, test cases và live deployment.",
):
    add_bullet(doc, item)

h = doc.add_heading("7. Rủi ro và phương án giảm thiểu", level=1)
keep_with_next(h)
risk_rows = [
    ("Scope quá rộng", "Cao", "Tối ưu lộ trình, bản đồ, memory và calendar có thể làm chậm luồng chính.", "Đóng băng MVP theo 8 acceptance criteria; phần nâng cao chỉ demo nếu P0 hoàn tất."),
    ("Code vẫn là template", "Cao", "Có cấu trúc nhưng chưa có domain logic nên dễ đánh giá sai tiến độ.", "Mỗi ngày phải có một artefact chạy được; không tính file placeholder là hoàn thành."),
    ("Tích hợp API ngoài", "Trung bình", "Calendar/map có thể lỗi quota, OAuth hoặc timeout.", "Thiết kế adapter + mock trước; luôn có timeout và trạng thái manual fallback."),
    ("Xung đột dữ liệu", "Cao", "Hai request đồng thời có thể cùng chốt một tài nguyên.", "Giữ exclusion constraint, row lock và transaction; thêm integration test cạnh tranh."),
    ("Rò rỉ secret/PII", "Cao", "Connection string Neon chứa mật khẩu; hội thoại chứa thông tin khách.", "Dùng .env/secret manager; không đưa secret vào frontend/Git; redact log và giới hạn quyền."),
    ("Thiếu owner rõ ràng", "Trung bình", "Công việc dàn trải và khó tích hợp cuối sprint.", "Chốt một owner cho từng stream và một người giữ end-to-end demo."),
]
add_status_table(
    doc,
    ["Rủi ro", "Mức", "Tác động", "Giảm thiểu"],
    risk_rows,
    [1800, 1000, 2800, 3760],
)

h = doc.add_heading("8. Nội dung cần mentor góp ý", level=1)
keep_with_next(h)
questions = (
    "Phạm vi MVP trên có đủ để chứng minh giá trị AI Agent, hay mentor yêu cầu tích hợp Google/Outlook Calendar thật ngay trong lần demo đầu?",
    "Nhóm có thể mock notification và calendar nhưng giữ inventory/booking bằng PostgreSQL thật không?",
    "Đối với HITL, mức tối thiểu nên là sale bấm duyệt trên UI hay có thể duyệt qua một endpoint/admin screen đơn giản?",
    "Nên ưu tiên hoàn thiện một luồng khách -> sale end-to-end hay làm đầy đủ authentication trước?",
    "Bộ eval tối thiểu ở giai đoạn MVP nên gồm booking success, tool-call accuracy và conflict prevention hay cần thêm tiêu chí hội thoại?",
    "Schema 17 bảng lõi + 7 bảng mở rộng có phù hợp, hay nên tách các bảng nâng cao sang migration sau để giảm phạm vi?",
)
for question in questions:
    add_bullet(doc, question)

h = doc.add_heading("9. Kịch bản trình bày trong buổi Mentor Duty (5-7 phút)", level=1)
keep_with_next(h)
agenda_rows = [
    ("0:00-0:45", "Bài toán", "Điều phối xem nhà qua chat dễ trùng lịch và bỏ lỡ khách."),
    ("0:45-2:00", "MVP", "Trình bày luồng khách yêu cầu -> slot -> sale duyệt -> appointment + hold."),
    ("2:00-3:30", "Bằng chứng", "Mở Neon/DBeaver ERD: 24 bảng; chỉ rõ approvals, appointments, property_holds."),
    ("3:30-4:30", "Trạng thái", "Nói thẳng database đã xong, application layer đang triển khai từ template."),
    ("4:30-5:30", "Sprint kế tiếp", "Chốt vertical slice end-to-end và owner từng stream."),
    ("5:30-7:00", "Xin góp ý", "Hỏi mentor về mức tích hợp thật, HITL tối thiểu và tiêu chí eval."),
]
add_status_table(doc, ["Thời lượng", "Phần", "Thông điệp"], agenda_rows, [1600, 1900, 5860])

add_callout(
    doc,
    "Câu kết đề xuất",
    "Nhóm đã hoàn thiện nền tảng dữ liệu và hiểu rõ các ràng buộc nghiệp vụ. Điều nhóm cần mentor xác nhận hôm nay là lát cắt MVP nào tạo giá trị rõ nhất để tập trung hoàn thiện end-to-end trong sprint tiếp theo.",
    fill=PALE_GREEN,
    label_color="2D6A36",
)

doc.add_page_break()

h = doc.add_heading("10. Biên bản ghi nhanh trong buổi mentor", level=1)
keep_with_next(h)
add_body(doc, "Dùng trang này để ghi quyết định ngay trong buổi họp; cập nhật lại WORKLOG.md và JOURNAL.md trong ngày.")

notes_rows = [
    ("Phạm vi MVP được mentor chốt", ""),
    ("Tích hợp bắt buộc", ""),
    ("Hạng mục được phép mock", ""),
    ("Tiêu chí đánh giá ưu tiên", ""),
    ("Rủi ro mentor lưu ý", ""),
    ("Quyết định kỹ thuật", ""),
    ("Owner và deadline", ""),
]
table = doc.add_table(rows=1, cols=2)
set_table_geometry(table, [2900, 6460])
set_table_borders(table)
for idx, header_text in enumerate(("Nội dung", "Ghi chú / quyết định")):
    set_cell_shading(table.rows[0].cells[idx], NAVY)
    p = table.rows[0].cells[idx].paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run(p.add_run(header_text), size=10, bold=True, color=WHITE)
repeat_table_header(table.rows[0])
prevent_row_split(table.rows[0])
for label, value in notes_rows:
    row = table.add_row()
    prevent_row_split(row)
    set_cell_shading(row.cells[0], LIGHT_GRAY)
    p0 = row.cells[0].paragraphs[0]
    set_run(p0.add_run(label), size=10, bold=True, color=DARK_BLUE)
    p1 = row.cells[1].paragraphs[0]
    set_run(p1.add_run(value or " "), size=10)
    for _ in range(2):
        blank = row.cells[1].add_paragraph(" ")
        blank.paragraph_format.space_after = Pt(0)
        set_run(blank.runs[0], size=10)

h = doc.add_heading("Checklist ngay sau buổi họp", level=2)
keep_with_next(h)
for item in (
    "Cập nhật phạm vi MVP và acceptance criteria theo góp ý mentor.",
    "Gán owner và deadline cụ thể cho từng stream.",
    "Tạo issue/task từ kế hoạch sprint, ưu tiên P0.",
    "Cập nhật WORKLOG.md, JOURNAL.md và ARCHITECTURE.md thay cho placeholder hiện tại.",
    "Chốt kịch bản demo end-to-end và ngày integration checkpoint.",
):
    add_bullet(doc, item)

# Document metadata.
doc.core_properties.title = "Báo cáo tiến độ Mentor Duty lần 2 - XHome VisitOps"
doc.core_properties.subject = "MVP và tiến độ dự án AI Agent đặt lịch xem bất động sản O2O"
doc.core_properties.author = "Nhóm P-046"
doc.core_properties.keywords = "AI20K, Mentor Duty, MVP, XHome VisitOps, PostgreSQL, LangGraph"

# Keep table rows intact where reasonable and avoid widows.
for paragraph in doc.paragraphs:
    p_pr = paragraph._p.get_or_add_pPr()
    widow = p_pr.find(qn("w:widowControl"))
    if widow is None:
        widow = OxmlElement("w:widowControl")
        p_pr.append(widow)

OUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(OUT)
print(OUT)
